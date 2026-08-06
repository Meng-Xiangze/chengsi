# -*- coding: utf-8 -*-
"""Edit tool — precise file editing with exact-text anchors.

For text files: replace, insert, delete, prepend, append.
For .docx: find paragraph by text, replace with formatting markers."""

import ast
import os
import re
from pathlib import Path
from typing import Any

from tools.base import BaseTool
from tools._hashline import revision, split_text, validate_anchor
from tools._spreadsheet import edit_csv, edit_xlsx
from core.process_utils import normalize_path, optional_import

_VALID_OPS = {
    "replace", "delete", "insert_before", "insert_after", "prepend", "append",
    "replace_range", "delete_range", "insert_before_anchor", "insert_after_anchor",
    "replace_symbol", "delete_symbol",
}
_OP_ALIASES = {"add_before": "insert_before", "add_after": "insert_after", "remove": "delete"}


class Edit(BaseTool):
    """Atomic text editing with hash, symbol, and exact-text anchors."""

    @property
    def tool_name(self) -> str:
        return "edit"

    @property
    def description(self) -> str:
        return (
            "Atomic file editing. Prefer replace_symbol for complete Python definitions and "
            "replace_range with LINE:HASH anchors for multiline code; use oldText for short unique text. "
            "Operations also include delete_range and anchor-based insertion. "
            "Aliases: remove=delete, add_before=insert_before, add_after=insert_after. "
            "For CSV/XLSX: oldText must match one complete cell and replace/delete edits are supported. "
            "For .docx: matches paragraph text; newText supports **bold** *italic* ^super^ _sub_ "
            "{size:N} {color:...} markers. "
            "Multiple edits applied atomically; overlapping edits rejected."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "path": {
                "type": "string",
                "description": "File path to edit.",
            },
            "edits": {
                "type": "array",
                "description": (
                    "Atomic edits. replace_symbol uses symbol; replace_range/delete_range use start/end "
                    "LINE:HASH anchors; insert_*_anchor uses anchor. oldText operations remain supported "
                    "for short unique text. op defaults to replace. "
                    "For CSV/XLSX: oldText matches one complete cell; use replace or delete. "
                    "For .docx: oldText matches paragraph text."
                ),
                "items": {
                    "type": "object",
                    "properties": {
                        "oldText": {
                            "type": "string",
                            "description": "Text to find (anchor). Must be unique. Not needed for prepend/append.",
                        },
                        "newText": {
                            "type": "string",
                            "description": "Replacement or inserted text. Not needed for delete.",
                        },
                        "op": {
                            "type": "string",
                            "enum": sorted(_VALID_OPS),
                            "description": "Operation. Default: replace.",
                        },
                        "start": {"type": "string", "description": "Inclusive start LINE:HASH anchor."},
                        "end": {"type": "string", "description": "Inclusive end LINE:HASH anchor."},
                        "anchor": {"type": "string", "description": "LINE:HASH insertion anchor."},
                        "symbol": {"type": "string", "description": "Qualified Python symbol, e.g. AgentRuntime.observe."},
                    },
                },
            },
            "revision": {
                "type": "string",
                "description": "Optional. If set, edit is rejected when file content hash differs (stale-file guard).",
            },
        }

    # ------------------------------------------------------------------ #
    #  Entry
    # ------------------------------------------------------------------ #

    def run(self, arguments: dict[str, Any]) -> str:
        args = arguments or {}
        raw_path = normalize_path(str(args.get("path", "")))
        raw_edits = args.get("edits", [])
        rev = str(args.get("revision", "")).strip() or None

        if not raw_path:
            return "Error: path is required."
        if not raw_edits or not isinstance(raw_edits, list):
            return "Error: edits must be a list of edit objects."

        path = Path(raw_path).resolve()
        if not path.exists():
            return f"Error: file not found: {path}"

        # Normalise operations
        edits: list[dict] = []
        for ei, e in enumerate(raw_edits):
            if not isinstance(e, dict):
                return f"Error: edits[{ei}] must be an object."
            op = str(e.get("op", "replace")).strip().lower()
            op = _OP_ALIASES.get(op, op)
            if op not in _VALID_OPS:
                return f"Error: edits[{ei}] invalid op {op!r}. Valid: {', '.join(sorted(_VALID_OPS))}."
            old = str(e.get("oldText", ""))
            new = str(e.get("newText", ""))
            if op in ("replace", "delete", "insert_before", "insert_after") and not old:
                return f"Error: edits[{ei}] ({op}) requires oldText as anchor."
            if op in ("replace_range", "delete_range") and (not e.get("start") or not e.get("end")):
                return f"Error: edits[{ei}] ({op}) requires start and end anchors."
            if op in ("insert_before_anchor", "insert_after_anchor") and not e.get("anchor"):
                return f"Error: edits[{ei}] ({op}) requires anchor."
            if op in ("replace_symbol", "delete_symbol") and not e.get("symbol"):
                return f"Error: edits[{ei}] ({op}) requires symbol."
            if op.startswith("delete"):
                new = ""
            if op not in ("delete", "delete_range", "delete_symbol") and not new:
                return f"Error: edits[{ei}] ({op}) requires newText."
            edits.append({
                "op": op, "oldText": old, "newText": new,
                "start": e.get("start"), "end": e.get("end"),
                "anchor": e.get("anchor"), "symbol": e.get("symbol"),
            })

        ext = path.suffix.lower()

        try:
            if ext == ".docx":
                return self._edit_docx(path, edits)
            if ext == ".csv":
                return edit_csv(path, edits)
            if ext == ".xlsx":
                return edit_xlsx(path, edits)
            return self._edit_text(path, edits, rev)
        except Exception as exc:
            return f"Error editing {path}: {exc}"

    def is_mutating(self, arguments: dict[str, Any]) -> bool:
        return True

    # ------------------------------------------------------------------ #
    #  Text-file engine
    # ------------------------------------------------------------------ #

    @staticmethod
    def _python_symbols(text: str, path: Path) -> list[tuple[str, tuple[int, int]]]:
        tree = ast.parse(text, filename=str(path))
        symbols: list[tuple[str, tuple[int, int]]] = []

        def visit(body: list[ast.stmt], prefix: str = "") -> None:
            for node in body:
                if isinstance(node, (ast.ClassDef, ast.FunctionDef, ast.AsyncFunctionDef)):
                    name = f"{prefix}.{node.name}" if prefix else node.name
                    decorator_lines = [item.lineno for item in getattr(node, "decorator_list", [])]
                    start = min([node.lineno, *decorator_lines])
                    symbols.append((name, (start, getattr(node, "end_lineno", node.lineno))))
                    if isinstance(node, ast.ClassDef):
                        visit(node.body, name)

        visit(tree.body)
        return symbols

    @staticmethod
    def _edit_text(path: Path, edits: list[dict], rev: str | None) -> str:
        raw = path.read_bytes()
        has_bom = raw.startswith(b"\xef\xbb\xbf")
        original = raw.decode("utf-8-sig" if has_bom else "utf-8")
        current_revision = revision(original)
        if rev and current_revision != rev:
            return (
                f"Error: stale revision {rev}; current revision is {current_revision}. "
                "Re-read the file and retry with current anchors."
            )

        raw_lines = original.splitlines(keepends=True)
        layout = split_text(original)
        starts: list[int] = []
        position = 0
        for raw_line in raw_lines:
            starts.append(position)
            position += len(raw_line)
        if layout.lines and len(starts) < len(layout.lines):
            starts.append(position)

        def line_span(start_index: int, end_index: int) -> tuple[int, int]:
            start_pos = starts[start_index]
            end_pos = starts[end_index + 1] if end_index + 1 < len(starts) else len(original)
            return start_pos, end_pos

        python_symbols = None
        spans: list[dict[str, Any]] = []
        for ei, edit in enumerate(edits):
            op = edit["op"]
            start_pos = end_pos = 0
            if op == "prepend":
                start_pos = end_pos = 0
            elif op == "append":
                start_pos = end_pos = len(original)
            elif op in ("replace", "delete", "insert_before", "insert_after"):
                old = edit["oldText"]
                count = original.count(old)
                if count != 1:
                    return f"Error: edits[{ei}] oldText matched {count} locations; expected exactly one."
                start_pos = original.index(old)
                end_pos = start_pos + len(old)
                if op == "insert_before":
                    end_pos = start_pos
                elif op == "insert_after":
                    start_pos = end_pos
            elif op in ("replace_range", "delete_range"):
                try:
                    start_line = validate_anchor(layout.lines, edit["start"])
                    end_line = validate_anchor(layout.lines, edit["end"])
                except ValueError as exc:
                    return f"Error: edits[{ei}] {exc}"
                if start_line > end_line:
                    return f"Error: edits[{ei}] start anchor is after end anchor."
                start_pos, end_pos = line_span(start_line, end_line)
            elif op in ("insert_before_anchor", "insert_after_anchor"):
                try:
                    line_index = validate_anchor(layout.lines, edit["anchor"])
                except ValueError as exc:
                    return f"Error: edits[{ei}] {exc}"
                line_start, line_end = line_span(line_index, line_index)
                start_pos = end_pos = line_start if op == "insert_before_anchor" else line_end
            elif op in ("replace_symbol", "delete_symbol"):
                if path.suffix.lower() != ".py":
                    return f"Error: edits[{ei}] {op} supports Python files only."
                try:
                    python_symbols = python_symbols or Edit._python_symbols(original, path)
                except SyntaxError as exc:
                    return f"Error: cannot locate Python symbols because the file does not parse: {exc}"
                requested = str(edit["symbol"])
                matches = [(name, value) for name, value in python_symbols if name == requested]
                if not matches and "." not in requested:
                    matches = [(name, value) for name, value in python_symbols if name.rsplit(".", 1)[-1] == requested]
                if len(matches) != 1:
                    return f"Error: edits[{ei}] symbol {requested!r} matched {len(matches)} definitions."
                start_line, end_line = matches[0][1]
                start_pos, end_pos = line_span(start_line - 1, end_line - 1)
            spans.append({"start": start_pos, "end": end_pos, "ei": ei})

        for i, first in enumerate(spans):
            for second in spans[i + 1:]:
                first_consumes = first["start"] != first["end"]
                second_consumes = second["start"] != second["end"]
                overlaps = first["start"] < second["end"] and second["start"] < first["end"]
                insertion_inside = (
                    first_consumes and first["start"] < second["start"] < first["end"]
                ) or (
                    second_consumes and second["start"] < first["start"] < second["end"]
                )
                if overlaps or insertion_inside or (not first_consumes and not second_consumes and first["start"] == second["start"]):
                    return f"Error: edits[{first['ei']}] and edits[{second['ei']}] overlap; merge them."

        text = original
        newline = layout.newline
        for span in sorted(spans, key=lambda item: item["start"], reverse=True):
            new = edits[span["ei"]]["newText"]
            if newline != "\n":
                new = new.replace("\r\n", "\n").replace("\r", "\n").replace("\n", newline)
            text = text[:span["start"]] + new + text[span["end"]:]

        if path.suffix.lower() == ".py":
            try:
                ast.parse(text, filename=str(path))
            except SyntaxError as exc:
                return f"Error: edit would produce invalid Python: {exc}. No changes written."

        encoded = text.encode("utf-8")
        if has_bom:
            encoded = b"\xef\xbb\xbf" + encoded
        temp_path = path.with_name(f".{path.name}.chengsi-edit-{os.getpid()}.tmp")
        try:
            temp_path.write_bytes(encoded)
            os.replace(temp_path, path)
        finally:
            if temp_path.exists():
                temp_path.unlink()
        old_lines = len(layout.lines)
        new_lines = len(split_text(text).lines)
        return (
            f"Edited {path}: {len(edits)} operation(s). Lines: {old_lines} -> {new_lines}; "
            f"rev={revision(text)}"
        )

    # ------------------------------------------------------------------ #
    #  DOCX engine
    # ------------------------------------------------------------------ #

    @staticmethod
    def _strip_markers(text: str) -> str:
        """Strip inline formatting markers: **bold** → bold, _sub_ → sub, etc."""
        # Order matters: ** before *, _ before others
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'\^(.+?)\^', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        return text

    @staticmethod
    def _body_idx(body, child) -> int:
        """Find the index of an lxml child element in the body. Returns -1 if not found."""
        for i, c in enumerate(body):
            if c is child:
                return i
        return -1

    @staticmethod
    def _edit_docx(path: Path, edits: list[dict]) -> str:
        try:
            docx = optional_import("docx", "python-docx")
            from docx import Document
            from tools.write import Write
            from tools.read import Read
            from docx.oxml.ns import qn
        except ImportError:
            return "DOCX editing requires python-docx. Install with:\n  pip install python-docx"

        doc = Document(str(path))

        # Build a map: formatted_text → paragraph_element
        # Use the same formatter as read tool so models see consistent text
        para_map: dict[str, object] = {}  # formatted_text → paragraph
        para_list: list = []  # (formatted_text, paragraph)
        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
            if tag == "p":
                formatted = Read._format_paragraph_xml(block, qn, -1)
                para_map[formatted.strip()] = block
                para_list.append((formatted.strip(), block))

        # Also build plain-text map for fallback
        plain_map: dict[str, object] = {}
        for para in doc.paragraphs:
            plain_map[para.text.strip()] = para._element

        found_count = 0
        for ei, edit in enumerate(edits):
            op = edit["op"]
            old = edit["oldText"].strip()
            new = edit["newText"]

            if op in ("prepend", "append"):
                # Insert paragraph at beginning or end
                if op == "prepend":
                    first = doc.element.body.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
                    )
                    p_elem = doc.element.body.makeelement(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p", {}
                    )
                    if first is not None:
                        idx = Edit._body_idx(doc.element.body, first)
                        doc.element.body.insert(idx, p_elem)
                    else:
                        doc.element.body.append(p_elem)
                    para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                    Write._add_formatted_run(para, new)
                else:
                    para = doc.add_paragraph()
                    Write._add_formatted_run(para, new)
                found_count += 1
                continue

            if op in ("replace", "delete", "insert_before", "insert_after"):
                # Try formatted-text match first (what model sees from read tool)
                # Then fall back to plain text (para.text)
                old_stripped = old.strip()
                matched_elem = para_map.get(old_stripped)
                if matched_elem is None:
                    matched_elem = plain_map.get(old_stripped)

                if matched_elem is None and op != "insert_before":
                    return f"Error: edits[{ei}] no paragraph matches: {old[:120]!r}"

                if op == "delete":
                    # Remove paragraph from document body
                    doc.element.body.remove(matched_elem)
                    found_count += 1
                elif op == "replace":
                    # Clear and rebuild — need the paragraph object, not just the element
                    # Find the python-docx Paragraph wrapping this element
                    target_para = None
                    for p in doc.paragraphs:
                        if p._element is matched_elem:
                            target_para = p
                            break
                    if target_para is not None:
                        target_para.clear()
                        Write._add_formatted_run(target_para, new)
                    else:
                        # Fallback: clear the element's runs manually
                        for r in matched_elem.findall(qn("w:r")):
                            matched_elem.remove(r)
                        # Add a temporary paragraph to get the run, then move it
                        temp_p = doc.add_paragraph()
                        Write._add_formatted_run(temp_p, new)
                        # Move runs from temp
                        for r in list(temp_p._element.findall(qn("w:r"))):
                            temp_p._element.remove(r)
                            matched_elem.append(r)
                        # Remove temp
                        doc.element.body.remove(temp_p._element)
                    found_count += 1
                elif op == "insert_before":
                    new_p = doc.element.body.makeelement(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p", {}
                    )
                    idx = Edit._body_idx(doc.element.body, matched_elem)
                    if idx < 0:
                        return f"Error: edits[{ei}] cannot find anchor paragraph in document."
                    doc.element.body.insert(idx, new_p)
                    # Add formatted content to the new paragraph
                    # Use python-docx paragraph wrapping the new element
                    temp_para = doc.add_paragraph()
                    Write._add_formatted_run(temp_para, new)
                    # Move runs
                    for r in list(temp_para._element.findall(qn("w:r"))):
                        temp_para._element.remove(r)
                        new_p.append(r)
                    doc.element.body.remove(temp_para._element)
                    found_count += 1
                elif op == "insert_after":
                    new_p = doc.element.body.makeelement(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p", {}
                    )
                    idx = Edit._body_idx(doc.element.body, matched_elem)
                    if idx < 0:
                        return f"Error: edits[{ei}] cannot find anchor paragraph in document."
                    doc.element.body.insert(idx + 1, new_p)
                    temp_para = doc.add_paragraph()
                    Write._add_formatted_run(temp_para, new)
                    for r in list(temp_para._element.findall(qn("w:r"))):
                        temp_para._element.remove(r)
                        new_p.append(r)
                    doc.element.body.remove(temp_para._element)
                    found_count += 1

        doc.save(str(path))
        return f"Edited {path}: {found_count} paragraph operation(s) applied."
