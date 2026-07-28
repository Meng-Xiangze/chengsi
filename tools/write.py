# -*- coding: utf-8 -*-
"""Write tool — create / overwrite files. Text files get raw content, .docx gets
Markdown-like formatting (bold, italic, lists, tables, images, fonts)."""

import os, re
from pathlib import Path
from typing import Any

from tools.base import BaseTool
from tools._spreadsheet import write_csv, write_xlsx


class Write(BaseTool):
    """Create or overwrite a file on disk."""

    @property
    def tool_name(self) -> str:
        return "write"

    @property
    def description(self) -> str:
        return (
            "Create a new file or overwrite an existing file. CSV content uses standard comma-separated rows; "
            "XLSX content uses CSV rows and optional [Sheet: name] sections. "
            "For .docx files, uses a Markdown-like syntax supporting **bold**, "
            "*italic*, ^superscript^, _subscript_, {size:N}, {color:...}, "
            "headings, lists, images, and tables."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "path": {
                "type": "string",
                "description": "File path to write to. Creates parent directories automatically.",
            },
            "content": {
                "type": "string",
                "description": (
                    "Complete file content. For CSV: comma-separated rows. For XLSX: CSV rows with optional "
                    "[Sheet: name] sections. For .docx: Markdown-like syntax — "
                    "**bold** *italic* ^super^ _sub_  {size:18}big{/size}  {color:red}red{/color}  "
                    "#heading  -bullet  1.numbered  |table|rows|  ![](image)"
                ),
            },
        }

    # ------------------------------------------------------------------ #
    #  Entry
    # ------------------------------------------------------------------ #

    def run(self, arguments: dict[str, Any]) -> str:
        args = arguments or {}
        raw_path = str(args.get("path", "")).strip()
        content = str(args.get("content", ""))

        if not raw_path:
            return "Error: path is required."

        path = Path(raw_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        ext = path.suffix.lower()

        try:
            if ext == ".docx":
                return self._write_docx(path, content)
            if ext == ".csv":
                return write_csv(path, content)
            if ext == ".xlsx":
                return write_xlsx(path, content)
            return self._write_text(path, content)
        except Exception as e:
            return f"Error writing {path}: {e}"

    # ------------------------------------------------------------------ #
    #  Write: text
    # ------------------------------------------------------------------ #

    @staticmethod
    def _write_text(path: Path, content: str) -> str:
        path.write_text(content, encoding="utf-8")
        size = os.path.getsize(str(path))
        return f"Written {path} ({size:,} bytes)"

    # ------------------------------------------------------------------ #
    #  Write: DOCX  (Markdown-like → formatted .docx)
    # ------------------------------------------------------------------ #

    @staticmethod
    def _write_docx(path: Path, content: str) -> str:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            return (
                "DOCX writing requires the python-docx package. Install with:\n"
                "  pip install python-docx"
            )

        doc = Document()
        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        lines = content.split("\n")
        image_count = 0
        table_rows: list[list[str]] = []

        i = 0
        while i < len(lines):
            stripped = lines[i].strip()
            i += 1

            # ── Accumulate table rows ──
            if "|" in stripped and not stripped.startswith(("#", "-", "*", "1.")):
                cells = [c.strip() for c in stripped.split("|")]
                cells = [c for c in cells if c]
                if len(cells) >= 2:
                    table_rows.append(cells)
                    continue

            # Flush accumulated table rows
            if table_rows:
                Write._flush_table(doc, table_rows)
                table_rows = []

            # Skip empty lines
            if not stripped:
                continue

            # ── Images ──
            img_match = re.match(r"^!\[(.*?)\]\((.+?)\)$|^!\((.+?)\)$", stripped)
            if img_match:
                alt = img_match.group(1) or ""
                img_path = img_match.group(2) or img_match.group(3) or ""
                img_path = img_path.strip()
                p = Write._resolve_image(img_path)
                if p:
                    try:
                        run = doc.add_paragraph().add_run()
                        run.add_picture(str(p), width=Inches(4.5))
                        if alt:
                            doc.add_paragraph(alt, style="Caption")
                        image_count += 1
                    except Exception:
                        doc.add_paragraph(f"[Image not found: {img_path}]")
                else:
                    doc.add_paragraph(f"[Image not found: {img_path}]")
                continue

            # ── Headings ──
            heading_match = re.match(r"^(#{1,3})\s+(.+)", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(Write._apply_inline_format(text), level=min(level, 3))
                continue

            # ── Bullet lists ──
            bullet_match = re.match(r"^[-*+]\s+(.+)", stripped)
            if bullet_match:
                text = bullet_match.group(1)
                p = doc.add_paragraph(style="List Bullet")
                Write._add_formatted_run(p, text)
                continue

            # ── Numbered lists ──
            num_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
            if num_match:
                text = num_match.group(2)
                p = doc.add_paragraph(style="List Number")
                Write._add_formatted_run(p, text)
                continue

            # ── Normal paragraph ──
            p = doc.add_paragraph()
            Write._add_formatted_run(p, stripped)

        # Flush any remaining table rows
        if table_rows:
            Write._flush_table(doc, table_rows)

        doc.save(str(path))
        size = os.path.getsize(str(path))
        return (
            f"Written {path} ({size:,} bytes)"
            + (f", {image_count} image(s)" if image_count else "")
        )

    @staticmethod
    def _flush_table(doc, rows: list[list[str]]):
        """Create a table from accumulated rows. First row = header."""
        if not rows:
            return
        from docx.shared import Inches
        ncols = max(len(r) for r in rows)
        # Pad short rows
        for r in rows:
            while len(r) < ncols:
                r.append("")
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.style = "Light Grid Accent 1"
        for ri, row_data in enumerate(rows):
            for ci, cell_text in enumerate(row_data):
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                if ri == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True

    # ── Formatting helpers ──────────────────────────────────────────

    @staticmethod
    def _add_formatted_run(paragraph, text: str):
        """Parse inline markers and add runs with formatting.

        Supported markers (can nest; outer wins for same property):
          **bold**   *italic*   ^superscript^   _subscript_
          {size:18}...   {color:#FF0000}...   {size:14,color:blue}...

        Styled spans close at the next same-type marker or end of text.
        """
        from docx.shared import Pt, RGBColor

        # Tokenize: styled spans, then bold/italic/super/sub markers
        style_re = r'\{(size|color):([^}]+)\}'
        inline_re = r'(\*\*.+?\*\*|\*.+?\*|\^.+?\^|_.+?_|[^{}*^_]+)'

        # State stack for nested styles
        stack: list[dict] = [{}]  # current style dict
        pos = 0
        buf: list[str] = []
        active_text = ""

        while pos < len(text):
            # Check for style span opener
            sm = re.match(style_re, text[pos:])
            if sm:
                # Flush any accumulated text with current style
                if active_text:
                    Write._emit_run(paragraph, active_text, stack[-1])
                    active_text = ""
                prop = sm.group(1)  # "size" or "color"
                val = sm.group(2)  # e.g. "14" or "14,color:blue"
                new_style = dict(stack[-1])  # inherit parent
                # First token is for prop, rest are key:value pairs
                parts = [p.strip() for p in val.split(",")]
                if prop == "size":
                    try:
                        new_style["size"] = Pt(float(parts[0]))
                    except (ValueError, IndexError):
                        pass
                elif prop == "color":
                    new_style["color"] = parts[0]
                # Remaining parts: key:value pairs or boolean flags
                for extra in parts[1:]:
                    if ":" in extra:
                        k, v = extra.split(":", 1)
                        k, v = k.strip(), v.strip()
                        if k == "size":
                            try:
                                new_style["size"] = Pt(float(v))
                            except ValueError:
                                pass
                        elif k == "color":
                            new_style["color"] = v
                    else:
                        # Boolean flag: **bold**, *italic*, etc. without value
                        flag = extra.strip().strip("*^_")
                        if "bold" in flag.lower():
                            new_style["bold"] = True
                        if "italic" in flag.lower() or "ital" in flag.lower():
                            new_style["italic"] = True
                        if "super" in flag.lower():
                            new_style["superscript"] = True
                        if "sub" in flag.lower():
                            new_style["subscript"] = True
                stack.append(new_style)
                pos += sm.end()
                continue

            # Check for style span closer (just the } character alone means close)
            # Actually, style spans close implicitly at next { marker or end.
            # Use explicit {/size} or {/color} to close.
            cm = re.match(r'\{/(size|color)\}', text[pos:])
            if cm:
                if active_text:
                    Write._emit_run(paragraph, active_text, stack[-1])
                    active_text = ""
                if len(stack) > 1:
                    stack.pop()
                pos += cm.end()
                continue

            # Check for inline formatting markers
            im = re.match(inline_re, text[pos:])
            if im:
                # Flush plain text before styled token
                if active_text:
                    Write._emit_run(paragraph, active_text, stack[-1])
                    active_text = ""
                token = im.group(0)
                if token.startswith("**") and token.endswith("**"):
                    inner = token[2:-2]
                    s = dict(stack[-1])
                    s["bold"] = True
                    Write._emit_run(paragraph, inner, s)
                elif token.startswith("*") and token.endswith("*"):
                    inner = token[1:-1]
                    s = dict(stack[-1])
                    s["italic"] = True
                    Write._emit_run(paragraph, inner, s)
                elif token.startswith("^") and token.endswith("^"):
                    inner = token[1:-1]
                    s = dict(stack[-1])
                    s["superscript"] = True
                    Write._emit_run(paragraph, inner, s)
                elif token.startswith("_") and token.endswith("_"):
                    inner = token[1:-1]
                    s = dict(stack[-1])
                    s["subscript"] = True
                    Write._emit_run(paragraph, inner, s)
                else:
                    active_text += token
                pos += im.end()
            else:
                # Literal character, skip the '{' that wasn't matched
                active_text += text[pos]
                pos += 1

        # Flush remaining text + close all styles
        if active_text or text and pos == 0:
            # Simple path fallback if nothing matched
            if not active_text and text:
                active_text = text
            Write._emit_run(paragraph, active_text, stack[-1])

    @staticmethod
    def _emit_run(paragraph, text: str, style: dict):
        """Create a single run with the given style dict applied."""
        from docx.shared import Pt, RGBColor
        run = paragraph.add_run(text)
        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
        if style.get("superscript"):
            run.font.superscript = True
        if style.get("subscript"):
            run.font.subscript = True
        if "size" in style:
            run.font.size = style["size"]
        if "color" in style:
            try:
                run.font.color.rgb = RGBColor.from_string(style["color"])
            except Exception:
                pass

    @staticmethod
    def _apply_inline_format(text: str) -> str:
        """Strip formatting markers for plain-text fields like headings."""
        return re.sub(r"(\*\*|\*|\^|_)(.+?)\1", r"\2", text)

    @staticmethod
    def _resolve_image(img_path: str) -> Path | None:
        """Resolve an image path (absolute or relative to cwd)."""
        p = Path(img_path)
        if p.is_absolute() and p.is_file():
            return p
        # Try relative to cwd
        cwd_p = Path.cwd() / img_path
        if cwd_p.is_file():
            return cwd_p
        # Try desktop
        desktop_p = Path(os.path.expanduser("~")) / "Desktop" / img_path
        if desktop_p.is_file():
            return desktop_p
        return None
